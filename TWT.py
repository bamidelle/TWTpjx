# project_x_full_v1.py
# Single-file Project X — Steps 1..5 combined
# Requirements (recommended): streamlit, sqlalchemy, pandas, scikit-learn, joblib, plotly, python-docx, shap
# Run: streamlit run project_x_full_v1.py

import os
import time
import traceback
from datetime import datetime, timedelta
from typing import Tuple, Optional, Dict, Any, List

import streamlit as st
import pandas as pd

# Defensive optional imports
try:
    import plotly.express as px
except Exception:
    px = None

try:
    import joblib
except Exception:
    joblib = None

# sklearn defensive imports
try:
    from sklearn.model_selection import train_test_split
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.metrics import (
        accuracy_score, precision_score, recall_score, f1_score,
        roc_auc_score, confusion_matrix, roc_curve
    )
    from sklearn.preprocessing import OneHotEncoder, StandardScaler
    from sklearn.compose import ColumnTransformer
    from sklearn.pipeline import Pipeline
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

# python-docx for report generation
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# shap optional
try:
    import shap
except Exception:
    shap = None

# matplotlib for PDP fallback
try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None

# SQLAlchemy
from sqlalchemy import create_engine, Column, Integer, String, Float, Boolean, DateTime, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

# ---------------------------
# CONFIG
# ---------------------------
DB_FILE = os.getenv("PROJECT_X_DB", "project_x_full_v1.db")
DATABASE_URL = f"sqlite:///{DB_FILE}"
MODEL_DIR = "models"
MODEL_PATH = os.path.join(MODEL_DIR, "lead_conversion_model_v1.joblib")
ENC_PATH = os.path.join(MODEL_DIR, "model_encoders.joblib")
os.makedirs(MODEL_DIR, exist_ok=True)

UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploaded_files")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

Base = declarative_base()
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine)

# ---------------------------
# ORM MODELS
# ---------------------------
class LeadStatus:
    NEW = "New"
    CONTACTED = "Contacted"
    INSPECTION_SCHEDULED = "Inspection Scheduled"
    INSPECTION_COMPLETED = "Inspection Completed"
    ESTIMATE_SUBMITTED = "Estimate Submitted"
    AWARDED = "Awarded"
    LOST = "Lost"
    ALL = [NEW, CONTACTED, INSPECTION_SCHEDULED, INSPECTION_COMPLETED, ESTIMATE_SUBMITTED, AWARDED, LOST]

class Lead(Base):
    __tablename__ = "leads"
    id = Column(Integer, primary_key=True)
    source = Column(String, default="Unknown")
    source_details = Column(String, nullable=True)
    contact_name = Column(String, nullable=True)
    contact_phone = Column(String, nullable=True)
    contact_email = Column(String, nullable=True)
    property_address = Column(String, nullable=True)
    damage_type = Column(String, nullable=True)
    assigned_to = Column(String, nullable=True)
    notes = Column(Text, nullable=True)
    estimated_value = Column(Float, nullable=True)
    status = Column(String, default=LeadStatus.NEW)
    created_at = Column(DateTime, default=datetime.utcnow)
    sla_hours = Column(Integer, default=24)
    sla_entered_at = Column(DateTime, nullable=True)

    contacted = Column(Boolean, default=False)
    inspection_scheduled = Column(Boolean, default=False)
    inspection_scheduled_at = Column(DateTime, nullable=True)
    inspection_completed = Column(Boolean, default=False)
    estimate_submitted = Column(Boolean, default=False)
    estimate_submitted_at = Column(DateTime, nullable=True)

    awarded_comment = Column(Text, nullable=True)
    awarded_date = Column(DateTime, nullable=True)
    awarded_invoice = Column(String, nullable=True)  # user wanted awarded_invoice storage in field 'b'

    lost_comment = Column(Text, nullable=True)
    lost_date = Column(DateTime, nullable=True)

    qualified = Column(Boolean, default=False)
    # predicted probability column will be added via migrations if not present

class Estimate(Base):
    __tablename__ = "estimates"
    id = Column(Integer, primary_key=True)
    lead_id = Column(Integer, nullable=False)
    amount = Column(Float, default=0.0)
    details = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    approved = Column(Boolean, default=False)
    lost = Column(Boolean, default=False)

# ---------------------------
# DB utilities
# ---------------------------
def create_tables_and_migrate():
    Base.metadata.create_all(bind=engine)
    # run lightweight migration to add predicted_prob if absent
    run_migration_add_predicted_prob()

def get_session():
    return SessionLocal()

def run_migration_add_predicted_prob():
    # Check if predicted_prob exists; if not, add column
    try:
        from sqlalchemy import inspect
        insp = inspect(engine)
        cols = [c['name'] for c in insp.get_columns('leads')]
        if 'predicted_prob' not in cols:
            with engine.connect() as conn:
                conn.execute("ALTER TABLE leads ADD COLUMN predicted_prob FLOAT")
    except Exception:
        # SQLite ALTER TABLE is limited but adding a column with default works for simple case
        pass

def add_lead(session, *, source="Unknown", source_details=None, contact_name=None, contact_phone=None,
             contact_email=None, property_address=None, damage_type=None, assigned_to=None, notes=None,
             estimated_value=None, sla_hours=24, qualified=False):
    lead = Lead(
        source=source,
        source_details=source_details,
        contact_name=contact_name,
        contact_phone=contact_phone,
        contact_email=contact_email,
        property_address=property_address,
        damage_type=damage_type,
        assigned_to=assigned_to,
        notes=notes,
        estimated_value=float(estimated_value) if estimated_value is not None else None,
        sla_hours=int(sla_hours),
        sla_entered_at=datetime.utcnow(),
        qualified=bool(qualified)
    )
    session.add(lead)
    session.commit()
    session.refresh(lead)
    return lead

def create_estimate(session, lead_id, amount, details=""):
    est = Estimate(lead_id=int(lead_id), amount=float(amount or 0.0), details=details)
    session.add(est)
    session.commit()
    session.refresh(est)
    # update lead flags
    lead = session.query(Lead).filter(Lead.id == lead_id).first()
    if lead:
        lead.estimate_submitted = True
        lead.estimate_submitted_at = datetime.utcnow()
        lead.status = LeadStatus.ESTIMATE_SUBMITTED
        session.add(lead); session.commit()
    return est

def leads_df(session):
    rows = session.query(Lead).all()
    data = []
    for r in rows:
        data.append({
            "id": r.id,
            "source": r.source,
            "source_details": r.source_details,
            "contact_name": r.contact_name,
            "contact_phone": r.contact_phone,
            "contact_email": r.contact_email,
            "property_address": r.property_address,
            "damage_type": r.damage_type,
            "assigned_to": r.assigned_to,
            "notes": r.notes,
            "estimated_value": float(r.estimated_value or 0.0),
            "status": r.status,
            "created_at": r.created_at,
            "sla_hours": r.sla_hours,
            "sla_entered_at": r.sla_entered_at or r.created_at,
            "contacted": bool(r.contacted),
            "inspection_scheduled": bool(r.inspection_scheduled),
            "inspection_scheduled_at": r.inspection_scheduled_at,
            "inspection_completed": bool(r.inspection_completed),
            "estimate_submitted": bool(r.estimate_submitted),
            "awarded_date": r.awarded_date,
            "awarded_invoice": r.awarded_invoice,
            "lost_date": r.lost_date,
            "qualified": bool(r.qualified),
            "predicted_prob": getattr(r, "predicted_prob", None)
        })
    df = pd.DataFrame(data)
    if df.empty:
        df = pd.DataFrame(columns=[
            "id","source","source_details","contact_name","contact_phone","contact_email",
            "property_address","damage_type","assigned_to","notes","estimated_value","status",
            "created_at","sla_hours","sla_entered_at","contacted","inspection_scheduled","inspection_scheduled_at",
            "inspection_completed","estimate_submitted","awarded_date","awarded_invoice","lost_date","qualified","predicted_prob"
        ])
    return df

def estimates_df(session):
    rows = session.query(Estimate).all()
    data = []
    for r in rows:
        data.append({
            "id": r.id,
            "lead_id": r.lead_id,
            "amount": r.amount,
            "details": r.details,
            "created_at": r.created_at,
            "approved": bool(r.approved),
            "lost": bool(r.lost)
        })
    return pd.DataFrame(data)

# ---------------------------
# Utilities
# ---------------------------
def save_uploaded_file(uploaded_file, prefix="file"):
    if uploaded_file is None:
        return None
    fname = f"{prefix}_{int(datetime.utcnow().timestamp())}_{uploaded_file.name}"
    path = os.path.join(UPLOAD_FOLDER, fname)
    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return path

def calculate_remaining_sla(sla_entered_at, sla_hours):
    try:
        if sla_entered_at is None:
            sla_entered_at = datetime.utcnow()
        if isinstance(sla_entered_at, str):
            sla_entered_at = datetime.fromisoformat(sla_entered_at)
        deadline = sla_entered_at + timedelta(hours=int(sla_hours or 24))
        remain = deadline - datetime.utcnow()
        return max(remain.total_seconds(), 0.0), (remain.total_seconds() <= 0)
    except Exception:
        return float("inf"), False

def compute_priority_for_lead_row(lead_row, weights, ml_prob=None):
    """Compute priority score (0..1) with optional ML probability component."""
    try:
        val = float(lead_row.get("estimated_value") or 0.0)
        baseline = float(weights.get("value_baseline", 5000.0))
        value_score = min(1.0, val / max(1.0, baseline))
    except Exception:
        value_score = 0.0

    try:
        sla_entered = lead_row.get("sla_entered_at") or lead_row.get("created_at")
        if sla_entered is None:
            time_left_h = 9999.0
        else:
            if isinstance(sla_entered, str):
                sla_entered = datetime.fromisoformat(sla_entered)
            deadline = sla_entered + timedelta(hours=int(lead_row.get("sla_hours") or 24))
            time_left_h = max((deadline - datetime.utcnow()).total_seconds() / 3600.0, 0.0)
    except Exception:
        time_left_h = 9999.0

    sla_score = max(0.0, (72.0 - min(time_left_h, 72.0)) / 72.0)
    contacted_flag = 0.0 if bool(lead_row.get("contacted")) else 1.0
    inspection_flag = 0.0 if bool(lead_row.get("inspection_scheduled")) else 1.0
    estimate_flag = 0.0 if bool(lead_row.get("estimate_submitted")) else 1.0

    urgency_component = (contacted_flag * weights.get("contacted_w", 0.6) +
                        inspection_flag * weights.get("inspection_w", 0.5) +
                        estimate_flag * weights.get("estimate_w", 0.5))

    total_weight = (weights.get("value_weight", 0.5) +
                    weights.get("sla_weight", 0.35) +
                    weights.get("urgency_weight", 0.15))
    if total_weight <= 0:
        total_weight = 1.0

    score = (value_score * weights.get("value_weight", 0.5) +
             sla_score * weights.get("sla_weight", 0.35) +
             urgency_component * weights.get("urgency_weight", 0.15)) / total_weight
    score = max(0.0, min(score, 1.0))

    # incorporate ML probability gently
    if ml_prob is not None:
        # mix score: 75% rule-based + 25% ML prob
        score = max(0.0, min(1.0, score * 0.75 + ml_prob * 0.25))

    return score

def predict_lead_probability_safe(model, X_row_df):
    """Given a fitted sklearn pipeline model and a single-row feature DF, return win prob or None."""
    if model is None or not SKLEARN_AVAILABLE or X_row_df is None or X_row_df.empty:
        return None
    try:
        if hasattr(model, "predict_proba"):
            return float(model.predict_proba(X_row_df)[:, 1][0])
        elif hasattr(model, "predict"):
            return float(model.predict(X_row_df)[0])
        else:
            return None
    except Exception:
        return None

# ---------------------------
# ML: feature builder, pipeline, train/save/load, persistence
# ---------------------------
def build_feature_df(df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.Series]:
    """Return X (features) and y (labels) for model training.
       Label: whether lead is AWARDED (1) vs not awarded (0).
       Features: estimated_value, qualified, sla_hours, contacted, inspection_scheduled, estimate_submitted,
                 damage_type, source
    """
    if df is None or df.empty:
        return pd.DataFrame(), pd.Series(dtype=int)
    d = df.copy()
    d["label_awarded"] = (d["status"] == LeadStatus.AWARDED).astype(int)
    features = [
        "estimated_value", "qualified", "sla_hours", "contacted",
        "inspection_scheduled", "estimate_submitted", "damage_type", "source"
    ]
    for c in features:
        if c not in d.columns:
            d[c] = 0
    X = d[features].copy()
    X["estimated_value"] = X["estimated_value"].fillna(0.0).astype(float)
    X["qualified"] = X["qualified"].astype(int)
    X["sla_hours"] = X["sla_hours"].fillna(24).astype(int)
    X["contacted"] = X["contacted"].astype(int)
    X["inspection_scheduled"] = X["inspection_scheduled"].astype(int)
    X["estimate_submitted"] = X["estimate_submitted"].astype(int)
    X["damage_type"] = X["damage_type"].fillna("unknown").astype(str)
    X["source"] = X["source"].fillna("unknown").astype(str)
    y = d["label_awarded"]
    return X, y

def create_sklearn_pipeline():
    if not SKLEARN_AVAILABLE:
        return None
    numeric_cols = ["estimated_value", "qualified", "sla_hours", "contacted", "inspection_scheduled", "estimate_submitted"]
    categorical_cols = ["damage_type", "source"]
    preproc = ColumnTransformer(transformers=[
        ("num", StandardScaler(), numeric_cols),
        ("cat", OneHotEncoder(handle_unknown="ignore", sparse=False), categorical_cols)
    ], remainder="drop")
    model = Pipeline(steps=[
        ("pre", preproc),
        ("clf", RandomForestClassifier(n_estimators=120, max_depth=6, random_state=42))
    ])
    return model

def save_model(model, path=MODEL_PATH):
    if joblib is None:
        return False, "joblib not installed"
    try:
        joblib.dump(model, path)
        return True, path
    except Exception as e:
        return False, str(e)

def load_model(path=MODEL_PATH):
    if joblib is None:
        return None, "joblib not installed"
    if not os.path.exists(path):
        return None, "path not found"
    try:
        m = joblib.load(path)
        return m, None
    except Exception as e:
        return None, str(e)

def persist_prediction_to_db(lead_id: int, prob: float):
    """Store predicted probability into lead.predicted_prob (migration must have created column)."""
    try:
        s = get_session()
        lead = s.query(Lead).filter(Lead.id == lead_id).first()
        if lead:
            # set attribute; attribute exists after migration, but SQLAlchemy mapping allows setting even if column absent won't persist
            setattr(lead, "predicted_prob", float(prob))
            s.add(lead); s.commit()
            return True
    except Exception:
        pass
    return False

# ---------------------------
# Confidence band helper
# ---------------------------
def confidence_band(prob: float, width: float = 0.08) -> Tuple[float, float]:
    low = max(0.0, prob - width)
    high = min(1.0, prob + width)
    return low, high

# ---------------------------
# CSS / UI
# ---------------------------
APP_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700&family=Comfortaa:wght@700&display=swap');
:root{
  --bg: #ffffff;
  --text: #0b1220;
  --muted: #6b7280;
  --primary-blue: #2563eb;
  --money-green: #22c55e;
  --accent-orange: #f97316;
  --danger: #ef4444;
  --card-radius: 12px;
}
body, .stApp { background: var(--bg); color: var(--text); font-family: 'Poppins', sans-serif; }
.header { font-family: 'Comfortaa', cursive; font-size:20px; font-weight:700; color:var(--text); padding:8px 0; }
.metric-card { border-radius: var(--card-radius); padding:16px; margin:8px; color:#fff; display:inline-block; vertical-align:top; box-shadow: 0 6px 16px rgba(16,24,40,0.06); transition: transform .18s ease, box-shadow .18s ease; }
.metric-card:hover { transform: translateY(-6px); box-shadow: 0 12px 28px rgba(16,24,40,0.12); }
.stage-card { background: #000; color: #fff; padding:12px; border-radius:10px; margin:6px; box-shadow: 0 6px 14px rgba(0,0,0,0.06); }
.progress-bar { width:100%; height:10px; background:#e6e6e6; border-radius:8px; overflow:hidden; margin-top:8px; }
.progress-fill { height:100%; border-radius:8px; transition: width .4s ease; }
.kpi-title { color: #ffffff; font-weight:600; font-size:13px; margin-bottom:6px; opacity:0.95; }
.kpi-value { font-weight:800; font-size:28px; color:#ffffff; }
.kpi-note { font-size:12px; color:rgba(255,255,255,0.9); margin-top:6px; }
.btn-animated { padding:10px 18px; border-radius:10px; border:none; cursor:pointer; font-weight:700; transition: transform .15s ease, box-shadow .15s ease; }
.btn-animated:hover { transform: translateY(-3px); box-shadow: 0 8px 20px rgba(0,0,0,0.12); }
"""

# ---------------------------
# APP START
# ---------------------------
create_tables_and_migrate()

st.set_page_config(page_title="Project X — Pipeline (Steps 1-5)", layout="wide", initial_sidebar_state="expanded")
st.markdown(f"<style>{APP_CSS}</style>", unsafe_allow_html=True)
st.markdown("<div class='header'>Project X — Sales & Conversion Tracker (Steps 1–5)</div>", unsafe_allow_html=True)

# ---------------------------
# Sidebar controls
# ---------------------------
with st.sidebar:
    st.header("Controls")
    page = st.radio("Go to", ["Leads / Capture", "Pipeline Board", "Analytics & SLA", "ML Model", "Evaluation Dashboard", "AI Recommendations", "Pipeline Report", "Exports"], index=1)
    st.markdown("---")
    if "weights" not in st.session_state:
        st.session_state.weights = {
            "value_weight": 0.5, "sla_weight": 0.35, "urgency_weight": 0.15,
            "contacted_w": 0.6, "inspection_w": 0.5, "estimate_w": 0.5, "value_baseline": 5000.0
        }
    st.markdown("### Priority weight tuning")
    st.session_state.weights["value_weight"] = st.slider("Estimate value weight", 0.0, 1.0, float(st.session_state.weights["value_weight"]), step=0.05)
    st.session_state.weights["sla_weight"] = st.slider("SLA urgency weight", 0.0, 1.0, float(st.session_state.weights["sla_weight"]), step=0.05)
    st.session_state.weights["urgency_weight"] = st.slider("Flags urgency weight", 0.0, 1.0, float(st.session_state.weights["urgency_weight"]), step=0.05)
    st.markdown("---")
    st.markdown("### Model (optional)")
    st.write("Model file path (joblib):")
    model_path_in = st.text_input("Model path", value=MODEL_PATH)
    if st.button("Load model"):
        if joblib and os.path.exists(model_path_in):
            try:
                m, err = load_model(model_path_in)
                if m:
                    st.session_state.lead_model = m
                    st.success("Model loaded (session).")
                else:
                    st.error(f"Failed to load model: {err}")
            except Exception as e:
                st.error(f"Failed to load model: {e}")
        else:
            st.info("Model path not found or joblib not available.")
    if st.button("Clear loaded model"):
        st.session_state.lead_model = None
        st.success("Model cleared from session.")
    st.markdown("---")
    if st.button("Add Demo Lead"):
        s = get_session()
        add_lead(
            s,
            source="Google Ads",
            source_details="gclid=demo",
            contact_name="Demo Customer",
            contact_phone="+15550000",
            contact_email="demo@example.com",
            property_address="100 Demo Ave",
            damage_type="water",
            assigned_to="Alex",
            estimated_value=4500,
            notes="Demo lead",
            sla_hours=24,
            qualified=True
        )
        st.success("Demo lead added")

# session model holder
if "lead_model" not in st.session_state:
    st.session_state.lead_model = None
if "model_features" not in st.session_state:
    st.session_state.model_features = None
if "model_metrics" not in st.session_state:
    st.session_state.model_metrics = None
if "last_train_time" not in st.session_state:
    st.session_state.last_train_time = None
if "user_role" not in st.session_state:
    st.session_state.user_role = "free"  # demo role

# ---------------------------
# Page: Leads / Capture
# ---------------------------
if page == "Leads / Capture":
    st.header("📇 Lead Capture")
    with st.form("lead_form", clear_on_submit=True):
        c1, c2 = st.columns(2)
        with c1:
            source = st.selectbox("Lead Source", ["Google Ads", "Organic Search", "Referral", "Phone", "Insurance", "Other"])
            source_details = st.text_input("Source details (UTM / notes)", placeholder="utm_source=google...")
            contact_name = st.text_input("Contact name", placeholder="John Doe")
            contact_phone = st.text_input("Contact phone", placeholder="+1-555-0123")
            contact_email = st.text_input("Contact email", placeholder="name@example.com")
        with c2:
            property_address = st.text_input("Property address", placeholder="123 Main St, City, State")
            damage_type = st.selectbox("Damage type", ["water", "fire", "mold", "contents", "reconstruction", "other"])
            assigned_to = st.text_input("Assigned to", placeholder="Estimator name")
            qualified_choice = st.selectbox("Is the Lead Qualified?", ["No", "Yes"], index=0)
            sla_hours = st.number_input("SLA hours (first response)", min_value=1, value=24, step=1)
        notes = st.text_area("Notes", placeholder="Additional context...")
        estimated_value = st.number_input("Estimated value (USD)", min_value=0.0, value=0.0, step=100.0)
        submitted = st.form_submit_button("Create Lead", help="Create a new lead")
        if submitted:
            s = get_session()
            lead = add_lead(
                s,
                source=source,
                source_details=source_details,
                contact_name=contact_name,
                contact_phone=contact_phone,
                contact_email=contact_email,
                property_address=property_address,
                damage_type=damage_type,
                assigned_to=assigned_to,
                notes=notes,
                sla_hours=int(sla_hours),
                qualified=True if qualified_choice == "Yes" else False,
                estimated_value=float(estimated_value or 0.0)
            )
            st.success(f"Lead created (ID: {lead.id})")

    st.markdown("---")
    st.subheader("Recent leads")
    s = get_session()
    df = leads_df(s)
    if df.empty:
        st.info("No leads yet. Create one above.")
    else:
        st.dataframe(df.sort_values("created_at", ascending=False).head(50))

# ---------------------------
# Page: Pipeline Board
# ---------------------------
elif page == "Pipeline Board":
    st.header("TOTAL LEAD PIPELINE — KEY PERFORMANCE INDICATOR")
    st.markdown("<em>High-level pipeline performance at a glance. Use filters and cards to drill into details.</em>", unsafe_allow_html=True)

    s = get_session()
    leads = s.query(Lead).order_by(Lead.created_at.desc()).all()
    if not leads:
        st.info("No leads yet. Create one from Lead Capture.")
    else:
        df = leads_df(s)
        weights = st.session_state.weights

        total_leads = len(df)
        qualified_leads = int(df[df["qualified"] == True].shape[0]) if not df.empty else 0

        sla_success_count = df.apply(lambda r: bool(r.get("contacted")), axis=1).sum() if not df.empty else 0
        sla_success_pct = (sla_success_count / total_leads * 100) if total_leads else 0.0

        qualification_pct = (qualified_leads / total_leads * 100) if total_leads else 0.0

        awarded_count = int(df[df["status"] == LeadStatus.AWARDED].shape[0]) if not df.empty else 0
        lost_count = int(df[df["status"] == LeadStatus.LOST].shape[0]) if not df.empty else 0
        closed = awarded_count + lost_count
        conversion_rate = (awarded_count / closed * 100) if closed else 0.0

        inspection_scheduled_count = int(df[df["inspection_scheduled"] == True].shape[0]) if not df.empty else 0
        inspection_pct = (inspection_scheduled_count / qualified_leads * 100) if qualified_leads else 0.0

        estimate_sent_count = int(df[df["estimate_submitted"] == True].shape[0]) if not df.empty else 0

        pipeline_job_value = float(df["estimated_value"].sum()) if not df.empty else 0.0

        active_leads = total_leads - (awarded_count + lost_count)

        # If model loaded in session, compute probabilities and attach to df
        model = st.session_state.lead_model
        if model is not None and SKLEARN_AVAILABLE:
            X_all, _ = build_feature_df(df)
            try:
                proba = model.predict_proba(X_all)[:, 1]
                df["win_prob"] = proba
            except Exception:
                df["win_prob"] = None
        else:
            df["win_prob"] = df.get("predicted_prob", None)

        KPI_ITEMS = [
            ("#2563eb", "Active Leads", f"{active_leads}", "Leads currently in pipeline"),
            ("#0ea5a4", "SLA Success", f"{sla_success_pct:.1f}%", "Leads contacted within SLA"),
            ("#a855f7", "Qualification Rate", f"{qualification_pct:.1f}%", "Leads marked qualified"),
            ("#f97316", "Conversion Rate", f"{conversion_rate:.1f}%", "Won / Closed"),
            ("#ef4444", "Inspections Booked", f"{inspection_pct:.1f}%", "Qualified → Scheduled"),
            ("#6d28d9", "Estimates Sent", f"{estimate_sent_count}", "Estimates submitted"),
            ("#22c55e", "Pipeline Job Value", f"${pipeline_job_value:,.0f}", "Total pipeline job value")
        ]

        st.markdown("<div style='display:flex; flex-wrap:wrap; gap:8px; align-items:stretch;'>", unsafe_allow_html=True)
        for color, title, value, note in KPI_ITEMS:
            st.markdown(f"""
                <div class="metric-card" style="background: linear-gradient(90deg, {color}, {color}); width:24%; min-width:200px;">
                    <div class="kpi-title">{title}</div>
                    <div class="kpi-value" style="font-family:'Poppins';">{value}</div>
                    <div class="kpi-note">{note}</div>
                </div>
            """, unsafe_allow_html=True)
        st.markdown(f"<div style='width:24%;min-width:200px;margin:8px;'></div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("---")

        # Pipeline stages - donut chart
        st.markdown("### Lead Pipeline Stages")
        st.markdown("<em>Distribution of leads across pipeline stages. Use this to spot stage drop-offs quickly.</em>", unsafe_allow_html=True)
        stage_colors = {
            LeadStatus.NEW: "#2563eb",
            LeadStatus.CONTACTED: "#eab308",
            LeadStatus.INSPECTION_SCHEDULED: "#f97316",
            LeadStatus.INSPECTION_COMPLETED: "#14b8a6",
            LeadStatus.ESTIMATE_SUBMITTED: "#a855f7",
            LeadStatus.AWARDED: "#22c55e",
            LeadStatus.LOST: "#ef4444"
        }
        stage_counts = df["status"].value_counts().reindex(LeadStatus.ALL, fill_value=0)
        pie_df = pd.DataFrame({"status": stage_counts.index, "count": stage_counts.values})
        if pie_df["count"].sum() == 0:
            st.info("No leads available to show pipeline stages.")
        else:
            if px:
                fig = px.pie(pie_df, names="status", values="count", hole=0.45, color="status",
                             color_discrete_map=stage_colors)
                fig.update_traces(textposition='inside', textinfo='percent+label')
                fig.update_layout(margin=dict(t=10, b=10), legend=dict(orientation="h", y=-0.15))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.table(pie_df)

        st.markdown("---")

        # Priority leads Top 5
        st.markdown("### TOP 5 PRIORITY LEADS")
        st.markdown("<em>Highest urgency leads by priority score (0–1). Address these first.</em>", unsafe_allow_html=True)
        priority_list = []
        for _, row in df.iterrows():
            try:
                ml_prob = float(row.get("win_prob")) if row.get("win_prob") is not None else None
            except Exception:
                ml_prob = None
            try:
                score = compute_priority_for_lead_row(row, weights, ml_prob=ml_prob)
            except Exception:
                score = 0.0
            sla_sec, overdue = calculate_remaining_sla(row.get("sla_entered_at") or row.get("created_at"), row.get("sla_hours"))
            time_left_h = sla_sec / 3600.0 if sla_sec not in (None, float("inf")) else 9999.0
            priority_list.append({
                "id": int(row["id"]),
                "contact_name": row.get("contact_name") or "No name",
                "estimated_value": float(row.get("estimated_value") or 0.0),
                "time_left_hours": time_left_h,
                "priority_score": score,
                "status": row.get("status"),
                "sla_overdue": overdue,
                "conversion_prob": ml_prob,
                "damage_type": row.get("damage_type", "Unknown")
            })
        pr_df = pd.DataFrame(priority_list).sort_values("priority_score", ascending=False)

        if pr_df.empty:
            st.info("No priority leads to display.")
        else:
            for _, r in pr_df.head(5).iterrows():
                score = r["priority_score"]
                status = r["status"]
                status_color = stage_colors.get(status, "#000000")
                if score >= 0.7:
                    priority_color = "#ef4444"
                    priority_label = "🔴 CRITICAL"
                elif score >= 0.45:
                    priority_color = "#f97316"
                    priority_label = "🟠 HIGH"
                else:
                    priority_color = "#22c55e"
                    priority_label = "🟢 NORMAL"
                if r["sla_overdue"]:
                    sla_html = f"<span style='color:#ef4444;font-weight:700;'>❗ OVERDUE</span>"
                else:
                    hours_left = int(r['time_left_hours'])
                    mins_left = int((r['time_left_hours'] * 60) % 60)
                    sla_html = f"<span style='color:#ef4444;font-weight:700;'>⏳ {hours_left}h {mins_left}m left</span>"
                conv_html = ""
                if r["conversion_prob"] is not None:
                    conv_pct = r["conversion_prob"] * 100
                    conv_color = "#22c55e" if conv_pct > 70 else ("#f97316" if conv_pct > 40 else "#ef4444")
                    conv_html = f"<span style='color:{conv_color};font-weight:600;margin-left:12px;'>📊 {conv_pct:.0f}% Win Prob</span>"
                st.markdown(f"""
                <div style="background: linear-gradient(180deg, rgba(0,0,0,0.04), rgba(0,0,0,0.02)); padding:12px; border-radius:12px; margin-bottom:10px;">
                  <div style="display:flex; justify-content:space-between; align-items:center;">
                    <div style="flex:1;">
                      <div style="margin-bottom:6px;">
                        <span style="color:{priority_color}; font-weight:800;">{priority_label}</span>
                        <span style="display:inline-block; padding:6px 12px; border-radius:18px; font-size:12px; font-weight:600; margin-left:8px; background:{status_color}22; color:{status_color};">{status}</span>
                      </div>
                      <div style="font-size:16px; font-weight:800; color:var(--text);">#{int(r['id'])} — {r['contact_name']}</div>
                      <div style="font-size:13px; color:var(--muted); margin-top:6px;">{r['damage_type'].title()} | Est: <span style='color:var(--money-green); font-weight:800;'>${r['estimated_value']:,.0f}</span></div>
                      <div style="font-size:13px; margin-top:8px; color:var(--muted);">{sla_html} {conv_html}</div>
                    </div>
                    <div style="text-align:right; padding-left:18px;">
                      <div style="font-size:28px; font-weight:900; color:{priority_color};">{r['priority_score']:.2f}</div>
                      <div style="font-size:11px; color:var(--muted); text-transform:uppercase;">Priority</div>
                    </div>
                  </div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # All Leads section with explanation
        st.markdown("### 📋 All Leads (expand a card to edit / change status)")
        st.markdown("<em>Expand a lead to edit details, change status, upload invoice when awarded, and create estimates.</em>", unsafe_allow_html=True)
        for lead in leads:
            est_val_display = f"${lead.estimated_value:,.0f}" if lead.estimated_value else "$0"
            card_title = f"#{lead.id} — {lead.contact_name or 'No name'} — {lead.damage_type or 'Unknown'} — {est_val_display}"
            with st.expander(card_title, expanded=False):
                colA, colB = st.columns([3, 1])
                with colA:
                    st.write(f"**Source:** {lead.source or '—'}   |   **Assigned:** {lead.assigned_to or '—'}")
                    st.write(f"**Address:** {lead.property_address or '—'}")
                    st.write(f"**Notes:** {lead.notes or '—'}")
                    st.write(f"**Created:** {lead.created_at.strftime('%Y-%m-%d %H:%M') if lead.created_at else '—'}")
                    # show predicted prob if exists
                    try:
                        if hasattr(lead, "predicted_prob") and lead.predicted_prob is not None:
                            st.write(f"**Predicted Win Prob:** {lead.predicted_prob*100:.1f}%")
                    except Exception:
                        pass
                with colB:
                    entered = lead.sla_entered_at or lead.created_at
                    if isinstance(entered, str):
                        try:
                            entered = datetime.fromisoformat(entered)
                        except:
                            entered = datetime.utcnow()
                    if entered is None:
                        entered = datetime.utcnow()
                    deadline = entered + timedelta(hours=(lead.sla_hours or 24))
                    remaining = deadline - datetime.utcnow()
                    if remaining.total_seconds() <= 0:
                        sla_status_html = "<div style='color:var(--danger);font-weight:700;'>❗ OVERDUE</div>"
                    else:
                        hours = int(remaining.total_seconds() // 3600)
                        mins = int((remaining.total_seconds() % 3600) // 60)
                        sla_status_html = f"<div style='color:var(--danger);font-weight:700;'>⏳ {hours}h {mins}m</div>"
                    st.markdown(f"<div style='text-align:right;'><div style='display:inline-block; padding:6px 12px; border-radius:20px; background:{stage_colors.get(lead.status,'#000')}22; color:{stage_colors.get(lead.status,'#000')}; font-weight:700;'>{lead.status}</div><div style='margin-top:12px;'>{sla_status_html}</div></div>", unsafe_allow_html=True)

                st.markdown("---")
                # Quick contact buttons
                qc1, qc2, qc3, qc4 = st.columns([1,1,1,4])
                phone = (lead.contact_phone or "").strip()
                email = (lead.contact_email or "").strip()
                if phone:
                    with qc1:
                        st.markdown(f"<a href='tel:{phone}'><button class='btn-animated' style='background:var(--primary-blue); color:#fff; border-radius:8px;'>📞 Call</button></a>", unsafe_allow_html=True)
                    with qc2:
                        wa_number = phone.lstrip("+").replace(" ", "").replace("-", "")
                        wa_link = f"https://wa.me/{wa_number}?text=Hi%2C%20following%20up%20on%20your%20restoration%20request."
                        st.markdown(f"<a href='{wa_link}' target='_blank'><button class='btn-animated' style='background:var(--money-green); color:#000; border-radius:8px;'>💬 WhatsApp</button></a>", unsafe_allow_html=True)
                else:
                    qc1.write(" "); qc2.write(" ")

                if email:
                    with qc3:
                        st.markdown(f"<a href='mailto:{email}?subject=Follow%20up'><button class='btn-animated' style='background:transparent; color:var(--text); border:1px solid #e5e7eb; border-radius:8px;'>✉️ Email</button></a>", unsafe_allow_html=True)
                else:
                    qc3.write(" ")
                qc4.write("")

                st.markdown("---")

                # Lead update form
                with st.form(f"update_lead_{lead.id}"):
                    st.markdown("#### Update Lead")
                    u1, u2 = st.columns(2)
                    with u1:
                        new_status = st.selectbox("Status", LeadStatus.ALL, index=LeadStatus.ALL.index(lead.status) if lead.status in LeadStatus.ALL else 0, key=f"status_{lead.id}")
                        new_assigned = st.text_input("Assigned to", value=lead.assigned_to or "", key=f"assign_{lead.id}")
                        new_contacted = st.checkbox("Contacted", value=bool(lead.contacted), key=f"contacted_{lead.id}")
                    with u2:
                        insp_sched = st.checkbox("Inspection Scheduled", value=bool(lead.inspection_scheduled), key=f"insp_sched_{lead.id}")
                        insp_comp = st.checkbox("Inspection Completed", value=bool(lead.inspection_completed), key=f"insp_comp_{lead.id}")
                        est_sub = st.checkbox("Estimate Submitted", value=bool(lead.estimate_submitted), key=f"est_sub_{lead.id}")
                    new_notes = st.text_area("Notes", value=lead.notes or "", key=f"notes_{lead.id}")
                    new_est_val = st.number_input("Job Value Estimate (USD)", value=float(lead.estimated_value or 0.0), min_value=0.0, step=100.0, key=f"estval_{lead.id}")

                    awarded_invoice_file = None
                    award_comment = None
                    lost_comment = None
                    if new_status == LeadStatus.AWARDED:
                        st.markdown("**Award details**")
                        award_comment = st.text_area("Award comment", key=f"award_comment_{lead.id}")
                        awarded_invoice_file = st.file_uploader("Upload Invoice File (optional) — only for Awarded", type=["pdf","jpg","jpeg","png","xlsx","csv"], key=f"award_inv_{lead.id}")
                    elif new_status == LeadStatus.LOST:
                        st.markdown("**Lost details**")
                        lost_comment = st.text_area("Lost comment", key=f"lost_comment_{lead.id}")

                    if st.form_submit_button("💾 Update Lead"):
                        try:
                            db = get_session()
                            db_lead = db.query(Lead).filter(Lead.id == lead.id).first()
                            if db_lead:
                                db_lead.status = new_status
                                db_lead.assigned_to = new_assigned
                                db_lead.contacted = bool(new_contacted)
                                db_lead.inspection_scheduled = bool(insp_sched)
                                db_lead.inspection_completed = bool(insp_comp)
                                db_lead.estimate_submitted = bool(est_sub)
                                db_lead.notes = new_notes
                                db_lead.estimated_value = float(new_est_val or 0.0)
                                if db_lead.sla_entered_at is None:
                                    db_lead.sla_entered_at = datetime.utcnow()
                                if new_status == LeadStatus.AWARDED:
                                    db_lead.awarded_date = datetime.utcnow()
                                    db_lead.awarded_comment = award_comment
                                    if awarded_invoice_file is not None:
                                        path = save_uploaded_file(awarded_invoice_file, prefix=f"lead_{db_lead.id}_inv")
                                        db_lead.awarded_invoice = path  # store invoice path in 'awarded_invoice'
                                if new_status == LeadStatus.LOST:
                                    db_lead.lost_date = datetime.utcnow()
                                    db_lead.lost_comment = lost_comment
                                db.add(db_lead)
                                db.commit()
                                st.success(f"Lead #{db_lead.id} updated.")
                                # Optionally recompute prediction for this lead if model exists
                                if st.session_state.lead_model is not None and SKLEARN_AVAILABLE:
                                    try:
                                        # rebuild X row for this lead
                                        df_single = leads_df(get_session())
                                        row = df_single[df_single["id"] == db_lead.id]
                                        X_row, _ = build_feature_df(row)
                                        if not X_row.empty:
                                            prob = predict_lead_probability_safe(st.session_state.lead_model, X_row)
                                            if prob is not None:
                                                persist_prediction_to_db(db_lead.id, prob)
                                    except Exception:
                                        pass
                            else:
                                st.error("Lead not found.")
                        except Exception as e:
                            st.error(f"Failed to update lead: {e}")
                            st.write(traceback.format_exc())

# ---------------------------
# Page: Analytics & SLA
# ---------------------------
elif page == "Analytics & SLA":
    st.header("📈 Analytics — CPA & Conversion Velocity (Date Range)")
    st.markdown("<em>Compare CPA per won job and conversion velocity over a selectable date range. Also see pipeline donut and SLA trends.</em>", unsafe_allow_html=True)

    s = get_session()
    df = leads_df(s)
    if df.empty:
        st.info("No leads to analyze. Add some leads first.")
    else:
        min_date = df["created_at"].min()
        max_date = df["created_at"].max()
        col_start, col_end = st.columns(2)
        start_date = col_start.date_input("Start date", min_value=min_date.date() if min_date is not None else datetime.utcnow().date(), value=min_date.date() if min_date is not None else datetime.utcnow().date())
        end_date = col_end.date_input("End date", min_value=start_date, value=max_date.date() if max_date is not None else datetime.utcnow().date())

        st.markdown("### Source spend mapping (for CPA calculation)")
        spend_input = st.text_input("Source spend mapping", value="", placeholder="Google Ads:1200,Referral:0")
        spend_map = {}
        if spend_input:
            try:
                parts = [p.strip() for p in spend_input.split(",") if p.strip()]
                for p in parts:
                    if ":" in p:
                        src, amt = p.split(":", 1)
                        spend_map[src.strip()] = float(amt.strip())
            except Exception:
                st.warning("Failed to parse spend input; please use format Source:amount, ...")

        start_dt = datetime.combine(start_date, datetime.min.time())
        end_dt = datetime.combine(end_date, datetime.max.time())
        df_range = df[(df["created_at"] >= start_dt) & (df["created_at"] <= end_dt)].copy()

        st.markdown("#### Pipeline Stages (donut)")
        stage_counts = df_range["status"].value_counts().reindex(LeadStatus.ALL, fill_value=0)
        pie_df = pd.DataFrame({"status": stage_counts.index, "count": stage_counts.values})
        if pie_df["count"].sum() == 0:
            st.info("No leads in selected range.")
        else:
            if px:
                fig = px.pie(pie_df, names="status", values="count", hole=0.45, color="status", color_discrete_map={
                    LeadStatus.NEW: "#2563eb", LeadStatus.CONTACTED: "#eab308", LeadStatus.INSPECTION_SCHEDULED: "#f97316",
                    LeadStatus.INSPECTION_COMPLETED: "#14b8a6", LeadStatus.ESTIMATE_SUBMITTED: "#a855f7", LeadStatus.AWARDED: "#22c55e",
                    LeadStatus.LOST: "#ef4444"
                })
                fig.update_traces(textposition='inside', textinfo='percent+label')
                fig.update_layout(margin=dict(t=10, b=10))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.table(pie_df)

        df_awarded = df_range[df_range["status"] == LeadStatus.AWARDED].copy()
        def parse_awarded_date(val, created):
            if pd.isna(val) or val is None:
                return created
            if isinstance(val, str):
                try:
                    return datetime.fromisoformat(val)
                except:
                    return created
            return val
        if not df_awarded.empty:
            df_awarded["awarded_date_parsed"] = df_awarded.apply(lambda r: parse_awarded_date(r.get("awarded_date"), r.get("created_at")), axis=1)
            df_awarded["month"] = df_awarded["awarded_date_parsed"].dt.to_period("M").astype(str)
            months = sorted(df_awarded["month"].unique().tolist())
            rows = []
            for m in months:
                month_df = df_awarded[df_awarded["month"] == m]
                won_count = len(month_df)
                total_spend = 0.0
                for src, amt in spend_map.items():
                    total_spend += amt
                cpa = (total_spend / won_count) if won_count and total_spend else (None if not total_spend else (total_spend / won_count))
                durations = []
                for _, rr in month_df.iterrows():
                    try:
                        created = rr["created_at"]
                        awarded = rr["awarded_date_parsed"]
                        if pd.isna(awarded) or awarded is None:
                            continue
                        durations.append((awarded - created).total_seconds() / 3600.0)
                    except:
                        continue
                avg_velocity = (sum(durations) / len(durations)) if durations else None
                rows.append({"month": m, "won_count": won_count, "total_spend": total_spend, "cpa": cpa, "avg_velocity_hrs": avg_velocity})
            stats_df = pd.DataFrame(rows)
            st.markdown("#### CPA per Won Job & Conversion Velocity (by month)")
            if px and not stats_df.empty:
                fig = px.line(stats_df, x="month", y=["cpa", "avg_velocity_hrs"], markers=True)
                fig.update_layout(yaxis_title="Value", xaxis_title="Month", legend_title="")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.dataframe(stats_df)

            st.markdown("**Notes:**")
            st.markdown("- **CPA per won job:** trending downward MoM, segmented by source. Provide accurate spend per source in the mapping above for precise CPA.")
            st.markdown("- **Velocity:** always improving; >48–72 hours stagnation = red flag lead.")
        else:
            st.info("No awarded jobs in selected date range to compute CPA/velocity.")

        # SLA Overdue Trend (last 30 days) and table
        st.markdown("---")
        st.subheader("SLA / Overdue Leads")
        st.markdown("<em>Trend of SLA overdue counts (last 30 days) and current overdue leads table.</em>", unsafe_allow_html=True)
        today = datetime.utcnow().date()
        days_back = 30
        ts_rows = []
        for d in range(days_back, -1, -1):
            day = today - pd.Timedelta(days=d)
            day_start = datetime.combine(day, datetime.min.time())
            day_end = datetime.combine(day, datetime.max.time())
            overdue_count = 0
            for _, row in df_range.iterrows():
                sla_entered = row.get("sla_entered_at") or row.get("created_at")
                try:
                    if pd.isna(sla_entered) or sla_entered is None:
                        sla_entered = row.get("created_at") or datetime.utcnow()
                    elif isinstance(sla_entered, str):
                        sla_entered = datetime.fromisoformat(sla_entered)
                except:
                    sla_entered = row.get("created_at") or datetime.utcnow()
                deadline = sla_entered + timedelta(hours=int(row.get("sla_hours") or 24))
                if deadline <= day_end and row.get("status") not in (LeadStatus.AWARDED, LeadStatus.LOST):
                    overdue_count += 1
            ts_rows.append({"date": day, "overdue_count": overdue_count})
        ts_df = pd.DataFrame(ts_rows)
        if not ts_df.empty:
            if px:
                fig = px.line(ts_df, x="date", y="overdue_count", markers=True, labels={"overdue_count": "Overdue leads"})
                fig.update_layout(margin=dict(t=6,b=6))
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.dataframe(ts_df)

        overdue_rows = []
        for _, row in df_range.iterrows():
            sla_entered = row.get("sla_entered_at") or row.get("created_at")
            try:
                if pd.isna(sla_entered) or sla_entered is None:
                    sla_entered = datetime.utcnow()
                elif isinstance(sla_entered, str):
                    sla_entered = datetime.fromisoformat(sla_entered)
            except:
                sla_entered = datetime.utcnow()
            sla_hours = int(row.get("sla_hours") or 24)
            deadline = sla_entered + timedelta(hours=sla_hours)
            overdue = deadline < datetime.utcnow() and row.get("status") not in (LeadStatus.AWARDED, LeadStatus.LOST)
            overdue_rows.append({
                "id": row.get("id"),
                "contact": row.get("contact_name"),
                "status": row.get("status"),
                "deadline": deadline,
                "overdue": overdue
            })
        df_overdue = pd.DataFrame(overdue_rows)
        if not df_overdue.empty:
            st.dataframe(df_overdue[df_overdue["overdue"] == True].sort_values("deadline"))
        else:
            st.info("No SLA overdue leads.")

# ---------------------------
# Page: ML Model (Train / Evaluate / Save / Auto-Retrain)
# ---------------------------
elif page == "ML Model":
    st.header("🧠 ML Model — Train, Evaluate & Save")
    st.markdown("<em>Build and evaluate a model to predict a lead's probability of converting to a won job.</em>", unsafe_allow_html=True)

    if not SKLEARN_AVAILABLE:
        st.error("Scikit-learn not available. Install scikit-learn to enable training and evaluation.")
    else:
        s = get_session()
        df = leads_df(s)
        if df.empty:
            st.info("No leads to train on. Collect labeled data (won jobs) first.")
        else:
            X, y = build_feature_df(df)
            st.markdown(f"Dataset size: **{len(X)}** leads — awarded (positive) count: **{int(y.sum())}**")
            test_size = st.slider("Test set size (%)", 5, 50, 20)
            random_state = st.number_input("Random seed", min_value=0, value=42, step=1)
            n_estimators = st.number_input("RandomForest n_estimators", min_value=10, value=120, step=10)
            max_depth = st.number_input("RandomForest max_depth", min_value=2, value=6, step=1)

            if st.button("Train model (RandomForest)"):
                try:
                    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size / 100.0, random_state=int(random_state), stratify=y if y.nunique()>1 else None)
                    model = create_sklearn_pipeline()
                    model.set_params(clf__n_estimators=int(n_estimators), clf__max_depth=int(max_depth))
                    model.fit(X_train, y_train)
                    # Evaluate
                    y_pred = model.predict(X_test)
                    y_proba = model.predict_proba(X_test)[:, 1] if hasattr(model, "predict_proba") else None
                    acc = accuracy_score(y_test, y_pred)
                    prec = precision_score(y_test, y_pred, zero_division=0)
                    rec = recall_score(y_test, y_pred, zero_division=0)
                    f1 = f1_score(y_test, y_pred, zero_division=0)
                    roc = roc_auc_score(y_test, y_proba) if y_proba is not None and len(set(y_test))>1 else None
                    cm = confusion_matrix(y_test, y_pred)
                    metrics = {"accuracy": acc, "precision": prec, "recall": rec, "f1": f1, "roc_auc": roc, "confusion_matrix": cm}

                    saved, p = save_model(model)
                    st.session_state.lead_model = model
                    st.session_state.model_features = list(X_train.columns)
                    st.session_state.model_metrics = metrics
                    st.session_state.last_train_time = time.time()
                    if saved:
                        st.success(f"Model trained and saved to {p}")
                    else:
                        st.warning(f"Model trained but not saved: {p}")
                    st.write(metrics)
                    if y_proba is not None and px:
                        fpr, tpr, _ = roc_curve(y_test, y_proba)
                        roc_df = pd.DataFrame({"fpr": fpr, "tpr": tpr})
                        fig = px.line(roc_df, x="fpr", y="tpr", title="ROC Curve (test set)")
                        fig.add_shape(type='line', x0=0, x1=1, y0=0, y1=1, line_dash='dash')
                        st.plotly_chart(fig, use_container_width=True)
                    st.markdown("Confusion matrix")
                    st.dataframe(pd.DataFrame(cm, index=["Actual 0","Actual 1"], columns=["Pred 0","Pred 1"]))
                except Exception as e:
                    st.error(f"Model training failed: {e}")
                    st.write(traceback.format_exc())

            st.markdown("---")
            st.markdown("### Auto-Retrain Controls")
            retrain_if_older_days = st.number_input("Retrain if older than (days)", min_value=1, value=7, step=1)
            do_check = st.checkbox("Run retrain-if-old check now (runs on page load only)", value=False)
            if do_check:
                last = st.session_state.get("last_train_time")
                now = time.time()
                if last is None or (now - last) > (retrain_if_older_days * 24 * 3600):
                    st.info("Training required — starting retrain now.")
                    try:
                        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size / 100.0, random_state=int(random_state), stratify=y if y.nunique()>1 else None)
                        model = create_sklearn_pipeline()
                        model.set_params(clf__n_estimators=int(n_estimators), clf__max_depth=int(max_depth))
                        model.fit(X_train, y_train)
                        y_pred = model.predict(X_test)
                        y_proba = model.predict_proba(X_test)[:, 1] if hasattr(model, "predict_proba") else None
                        acc = accuracy_score(y_test, y_pred)
                        roc = roc_auc_score(y_test, y_proba) if y_proba is not None and len(set(y_test))>1 else None
                        cm = confusion_matrix(y_test, y_pred)
                        metrics = {"accuracy": acc, "roc_auc": roc, "confusion_matrix": cm}
                        saved, p = save_model(model)
                        st.session_state.lead_model = model
                        st.session_state.model_features = list(X_train.columns)
                        st.session_state.model_metrics = metrics
                        st.session_state.last_train_time = time.time()
                        st.success("Auto-retrain done and model saved.")
                    except Exception as e:
                        st.error(f"Auto-retrain failed: {e}")
                        st.write(traceback.format_exc())
                else:
                    st.info("Recent training detected — no retrain necessary.")

# ---------------------------
# Page: Evaluation Dashboard
# ---------------------------
elif page == "Evaluation Dashboard":
    st.header("📊 Evaluation Dashboard")
    st.markdown("<em>Model performance metrics, confusion matrix, and ROC if available.</em>", unsafe_allow_html=True)
    model = st.session_state.get("lead_model")
    metrics = st.session_state.get("model_metrics")
    if model is None:
        st.info("No model loaded. Train one in 'ML Model' or load from sidebar.")
    else:
        st.markdown("### Stored model metrics")
        st.write(metrics or "No metrics stored.")
        if metrics and "confusion_matrix" in metrics and metrics["confusion_matrix"] is not None:
            cm = metrics["confusion_matrix"]
            st.markdown("#### Confusion Matrix")
            st.dataframe(pd.DataFrame(cm, index=["Actual 0","Actual 1"], columns=["Pred 0","Pred 1"]))
        if metrics and metrics.get("roc_auc") is not None:
            st.markdown(f"**ROC AUC:** {metrics['roc_auc']:.3f}")
        if st.button("Run quick evaluation on current DB"):
            try:
                s = get_session()
                df = leads_df(s)
                X, y = build_feature_df(df)
                if X.empty:
                    st.info("No data to evaluate.")
                else:
                    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y if y.nunique()>1 else None)
                    y_pred = model.predict(X_test)
                    y_proba = model.predict_proba(X_test)[:,1] if hasattr(model,"predict_proba") else None
                    st.write("Accuracy:", accuracy_score(y_test, y_pred))
                    st.write("Precision:", precision_score(y_test, y_pred, zero_division=0))
                    st.write("Recall:", recall_score(y_test, y_pred, zero_division=0))
                    st.write("F1:", f1_score(y_test, y_pred, zero_division=0))
                    if y_proba is not None and px:
                        fpr,tpr,_ = roc_curve(y_test, y_proba)
                        roc_df = pd.DataFrame({"fpr":fpr,"tpr":tpr})
                        fig = px.line(roc_df, x="fpr", y="tpr", title="ROC Curve (current DB)")
                        fig.add_shape(type='line', x0=0, x1=1, y0=0, y1=1, line_dash='dash')
                        st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.error(f"Evaluation failed: {e}")
                st.write(traceback.format_exc())

# ---------------------------
# Page: AI Recommendations
# ---------------------------
elif page == "AI Recommendations":
    st.header("💡 AI Recommendations")
    st.markdown("<em>Source-level recommendations and improvement notes.</em>", unsafe_allow_html=True)
    s = get_session()
    df = leads_df(s)
    if df.empty:
        st.info("No leads yet.")
    else:
        src_perf = df.groupby("source").apply(lambda d: (d["status"] == LeadStatus.AWARDED).mean()).reset_index()
        src_perf.columns = ["source", "win_rate"]
        st.markdown("### Source performance (win rate)")
        st.dataframe(src_perf.sort_values("win_rate", ascending=False))
        st.markdown("### Suggestions")
        best_sources = src_perf.sort_values("win_rate", ascending=False).head(3)
        worst_sources = src_perf.sort_values("win_rate", ascending=True).head(3)
        st.markdown("Top performing sources — consider increasing spend or focusing on these:")
        for _, r in best_sources.iterrows():
            st.write(f"- {r['source']}: win rate {r['win_rate']*100:.1f}%")
        st.markdown("Low performing sources — review targeting or lead quality:")
        for _, r in worst_sources.iterrows():
            st.write(f"- {r['source']}: win rate {r['win_rate']*100:.1f}%")

# ---------------------------
# Page: Pipeline Report (DOCX)
# ---------------------------
elif page == "Pipeline Report":
    st.header("📄 Pipeline Report (Export .docx)")
    st.markdown("<em>Generate a Word (.docx) pipeline report containing KPI summary, ML metrics and top leads.</em>", unsafe_allow_html=True)
    s = get_session()
    df = leads_df(s)
    if df.empty:
        st.info("No leads to report.")
    else:
        if not DOCX_AVAILABLE:
            st.warning("python-docx not installed — install it (pip install python-docx) to enable DOCX export.")
        else:
            if st.button("Generate pipeline_report.docx"):
                try:
                    doc = Document()
                    doc.add_heading('TOTAL LEAD PIPELINE KEY PERFORMANCE INDICATOR', 0)
                    doc.add_paragraph('*Total health view of live leads, response performance, and projected job value.*')
                    total_leads = len(df)
                    awarded_count = int(df[df["status"] == LeadStatus.AWARDED].shape[0])
                    pipeline_value = float(df["estimated_value"].sum()) if not df.empty else 0.0
                    doc.add_paragraph(f"Total leads: {total_leads}")
                    doc.add_paragraph(f"Awarded (won): {awarded_count}")
                    doc.add_paragraph(f"Pipeline job value: ${pipeline_value:,.0f}")
                    # ML metrics
                    metrics = st.session_state.get("model_metrics")
                    if metrics:
                        doc.add_heading("ML Metrics", level=1)
                        doc.add_paragraph(f"Accuracy: {metrics.get('accuracy')}")
                        doc.add_paragraph(f"ROC AUC: {metrics.get('roc_auc')}")
                    # Top 5 priority leads
                    # compute priority with available model prob
                    df_local = df.copy()
                    if st.session_state.lead_model is not None and SKLEARN_AVAILABLE:
                        try:
                            X_all, _ = build_feature_df(df_local)
                            proba = st.session_state.lead_model.predict_proba(X_all)[:, 1]
                            df_local["win_prob"] = proba
                        except Exception:
                            df_local["win_prob"] = None
                    priority_list = []
                    for _, row in df_local.iterrows():
                        ml_prob = float(row.get("win_prob")) if row.get("win_prob") is not None else None
                        score = compute_priority_for_lead_row(row, st.session_state.weights, ml_prob=ml_prob)
                        priority_list.append((int(row["id"]), row.get("contact_name") or "No name", score))
                    top5 = sorted(priority_list, key=lambda x: x[2], reverse=True)[:5]
                    doc.add_heading("Top 5 Priority Leads", level=1)
                    for pid, name, score in top5:
                        doc.add_paragraph(f"#{pid} — {name} — Score: {score:.2f}")
                    path = os.path.join(MODEL_DIR, f"pipeline_report_{int(time.time())}.docx")
                    doc.save(path)
                    with open(path, "rb") as fh:
                        st.download_button("Download pipeline_report.docx", fh, file_name=os.path.basename(path))
                    st.success("Report generated.")
                except Exception as e:
                    st.error(f"Failed to generate report: {e}")
                    st.write(traceback.format_exc())

# ---------------------------
# Page: Exports
# ---------------------------
elif page == "Exports":
    st.header("📤 Export data")
    s = get_session()
    df_leads = leads_df(s)
    if df_leads.empty:
        st.info("No leads yet to export.")
    else:
        csv = df_leads.to_csv(index=False).encode("utf-8")
        st.download_button("Download leads.csv", csv, file_name="leads.csv", mime="text/csv")
    df_est = estimates_df(s)
    if not df_est.empty:
        st.download_button("Download estimates.csv", df_est.to_csv(index=False).encode("utf-8"), file_name="estimates.csv", mime="text/csv")

# ---------------------------
# End
# ---------------------------
